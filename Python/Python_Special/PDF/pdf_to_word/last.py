import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches


def extract_pdf_content(pdf_path):
    doc = fitz.open(pdf_path)

    # Create a Word document
    word_doc = Document()

    for page_number in range(doc.page_count):
        page = doc[page_number]

        # Extract text
        text = page.get_text()
        word_doc.add_paragraph(text)

        # Extract images
        for img_index in range(page.get_image_count()):
            img = page.get_image(img_index)
            img_bytes = img.get_pixmap().tobytes()
            img_path = f'image_{page_number}_{img_index}.png'
            with open(img_path, 'wb') as img_file:
                img_file.write(img_bytes)
            word_doc.add_picture(img_path, width=Inches(2))  # Adjust width as needed

        # Extract tables (this is a basic example, may need further adjustments)
        for table_index, table in enumerate(page.get_tables()):
            word_table = word_doc.add_table(rows=len(table), cols=max(len(row) for row in table))
            for row_index, row in enumerate(table):
                for col_index, cell in enumerate(row):
                    word_table.cell(row_index, col_index).text = cell.get_text()

    doc.close()
    return word_doc


def save_word_document(word_doc, output_path):
    word_doc.save(output_path)


# Example usage
pdf_path = 'example.pdf'
word_document = extract_pdf_content(pdf_path)
output_path = 'output.docx'
save_word_document(word_document, output_path)