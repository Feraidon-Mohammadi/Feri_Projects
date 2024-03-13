import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


# Create a new Word document
doc = docx.Document()

# Define the elements
empfenger_info = {
    "name": "Feri M",
    "address": "Ritht Mianyer Str. 2202",
    "plz": "33303 Köln",
}

absender_info = {
    "name": "DD GmbH",
    "address": "Rocky Str. 22",
    "plz": "44055 Dortmund",
}



datum = datetime.date
# Add "Datum" section aligned to the right
datum_section = doc.add_paragraph()
datum_run = datum_section.add_run("den: " + datum)
datum_run.alignment = WD_ALIGN_PARAGRAPH.RIGHT


betreff_text = "Praktikum Stelle als Anwendungsentwicklung"

anrede = "Damen und Herren"

text_data = [
    "Hiermit bewerbe ich mich bei Ihnen um eine Praktikum Stelle im Bereich Anwendungsentwicklung.",
    "Durch meine bisherige Ausbildung und Erfahrung habe ich fundierte Kenntnisse in verschiedenen Programmiersprachen und Anwendungsentwicklung erworben.",
    "Ich bin motiviert und freue mich darauf, mein Wissen in der Praxis anzuwenden und weiter zu vertiefen.",
    "Mit freundlichen Grüßen,",
    "Feri M",
]

# Add "Empfänger" section
empfenger_section = doc.add_paragraph()
empfenger_section.add_run("Empfänger").bold = True
empfenger_section.alignment = WD_ALIGN_PARAGRAPH.LEFT

empfenger_name = doc.add_paragraph()
empfenger_name.add_run(empfenger_info["name"])
empfenger_name.add_run("\n" + empfenger_info["address"])
empfenger_name.add_run("\n" + empfenger_info["plz"])

# Add "Absender" section
absender_section = doc.add_paragraph()
absender_section.add_run("Absender").bold = True
absender_section.alignment = WD_ALIGN_PARAGRAPH.LEFT

absender_name = doc.add_paragraph()
absender_name.add_run(absender_info["name"])
absender_name.add_run("\n" + absender_info["address"])
absender_name.add_run("\n" + absender_info["plz"])

# Add "Betreff" section with font size 16
betreff_section = doc.add_paragraph()
betreff_run = betreff_section.add_run(betreff_text)
betreff_run.font.size = Pt(16)
betreff_section.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Add "Datum" section
datum_section = doc.add_paragraph()
datum_section.add_run("den," + datum)
datum_section.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Add "Anrede" section
anrede_section = doc.add_paragraph(anrede)
anrede_section.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Add text data with paragraphs
for paragraph_text in text_data:
    text_paragraph = doc.add_paragraph(paragraph_text)

# Save the document to a file with a .docx extension
doc.save("job_application.docx")