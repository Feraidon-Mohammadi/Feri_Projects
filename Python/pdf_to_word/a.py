import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def extract_pdf_content(pdf_path):
    doc = fitz.open(pdf_path)
    word_doc = Document()

    for page_number in range(doc.page_count):
        page = doc[page_number]

        # Extract text
        text = page.get_text()
        word_doc.add_paragraph(text)

        # Extract images
        for img_index, img_info in enumerate(page.get_images(full=True)):
            img_index += 1  # Image indexing starts from 1
            img = doc.extract_image(img_index)
            img_bytes = img["image"]

            # You may need to handle image data appropriately, such as saving it to a file or adding it to the Word document.
            # For simplicity, let's add a placeholder text to indicate the presence of an image.
            word_doc.add_paragraph(f"Image {img_index} - Image data: {len(img_bytes)} bytes")

    doc.close()
    return word_doc

def save_word_document(word_doc, output_path):
    word_doc.save(output_path)

# Example usage
pdf_path = 'example.pdf'
output_path = 'output.docx'

# Extract content from PDF
word_document = extract_pdf_content(pdf_path)

# Save Word document
save_word_document(word_document, output_path)